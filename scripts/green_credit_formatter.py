#!/usr/bin/env python3
"""
绿色信贷论文格式转换器
=====================
将 Markdown 论文转换为 LaTeX 和 Word 格式，
支持《金融研究》/《数量经济技术经济研究》排版规范。

使用方法：
  python scripts/green_credit_formatter.py              # 转换全部格式
  python scripts/green_credit_formatter.py --format latex  # 仅LaTeX
  python scripts/green_credit_formatter.py --format word    # 仅Word
"""

import argparse
from datetime import datetime
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent.parent

# ════════════════════════════════════════════════════════════════════
# LaTeX 模板
# ════════════════════════════════════════════════════════════════════

LATEX_TEMPLATE = r"""
% ============================================================
% 绿色信贷政策对重污染企业融资约束的影响研究
% ============================================================
% 格式：《金融研究》/ 数量经济技术经济研究
% 编译：XeLaTeX
% ============================================================

\documentclass[12pt,a4paper]{{article}}

% 页面设置
\usepackage[top=2.54cm,bottom=2.54cm,left=3.18cm,right=3.18cm]{{geometry}}
\usepackage{{setspace}}
\setstretch{{1.5}}

% 数学字体
\usepackage{{amsmath,amssymb}}
\usepackage{{bm}}

% 表格
\usepackage{{booktabs}}
\usepackage{{multirow}}
\usepackage{{threeparttable}}
\usepackage{{caption}}
\captionsetup{{font=normalsize, labelsep=quad}}

% 图表
\usepackage{{graphicx}}
\usepackage{{subfig}}

% 中文支持
\usepackage{{ctex}}
\setCJKmainfont{{[SimSun.ttc]}}[BoldFont={SimHei.ttc}, ItalicFont={KaiTi.ttf}]

% 参考文献
\usepackage[numbers]{{natbib}}
\bibliographystyle{{economic}}

% 超链接
\usepackage{{hyperref}}
\hypersetup{{
    colorlinks=true,
    linkcolor=blue,
    citecolor=blue,
    urlcolor=blue,
}}

\begin{{document}}

% ============================================================
% 标题页
% ============================================================
\begin{{titlepage}}
\vspace{{2cm}}
\begin{{center}}
{\LARGE\bfseries 绿色信贷政策对重污染企业融资约束的影响研究}\\[1.5cm]
\end{{center}}
\vspace{{1cm}}
\end{{titlepage}}

% ============================================================
% 摘要
% ============================================================
\begin{{abstract}}
绿色信贷作为中国绿色金融体系的核心政策工具，自2012年《绿色信贷指南》发布以来，
对高污染高排放企业的融资约束产生了深远影响。
本文以2012年银监会颁布《绿色信贷指南》作为准自然实验，
采用双重差分法（DID）和倾向得分匹配（PSM-DID）方法，
利用2008—2020年中国A股上市公司数据，
系统考察绿色信贷政策对重污染企业融资约束的影响及其作用机制。
研究发现：（1）绿色信贷政策显著加剧了重污染企业的融资约束，
具体表现为长期借款规模明显下降，而短期借款规模则呈现扩张态势；
（2）异质性分析表明，政策对民营企业融资约束的抑制效应显著强于国有企业，
存在明显的"所有制歧视"特征；
（3）机制检验发现，绿色信贷政策通过提高银行风险规避偏好和
加剧信息不对称两条路径约束了重污染企业的外部融资能力。

\textbf{{关键词}}：绿色信贷；融资约束；重污染企业；双重差分法；所有制歧视
\end{{abstract}}

\vspace{{1ex}}
\textbf{{JEL}}：G21；Q56；O16

\newpage

% ============================================================
% 正文
% ============================================================
%(CONTENT)s

% ============================================================
% 参考文献
% ============================================================
\newpage
\bibliography{green_credit_refs}

\end{{document}}
"""


# ════════════════════════════════════════════════════════════════════
# Markdown → LaTeX 转换器
# ════════════════════════════════════════════════════════════════════

def md_to_latex(text: str) -> str:
    """将 Markdown 转换为 LaTeX 代码"""
    lines = text.split("\n")
    result = []
    in_table = False
    table_lines = []
    in_math = False

    for line in lines:
        stripped = line.strip()

        # 跳过分隔线
        if stripped.startswith("---"):
            result.append("\\newpage")
            continue

        # 标题层级
        if stripped.startswith("# "):
            title = stripped[2:]
            result.append(f"\\section{{{title}}}")
        elif stripped.startswith("## "):
            title = stripped[3:]
            result.append(f"\\subsection{{{title}}}")
        elif stripped.startswith("### "):
            title = stripped[4:]
            result.append(f"\\subsubsection{{{title}}}")

        # 段落（跳过空行）
        elif stripped and not stripped.startswith("|") and not stripped.startswith("!"):
            if stripped.startswith("**") and stripped.endswith("**"):
                result.append(f"\\paragraph{{{stripped[2:-2]}}}")
            else:
                result.append(stripped + "\n")

        # 表格（转为 booktabs 格式）
        elif stripped.startswith("|"):
            if "---" in stripped:
                in_table = True
                continue
            # 清理表格单元格
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if in_table:
                table_lines.append(cells)
            else:
                # 第一个表头行
                table_lines = [cells]
                in_table = True
        else:
            if in_table and table_lines:
                # 输出 LaTeX 表格
                if len(table_lines) == 1:
                    header = " & ".join(table_lines[0])
                    result.append("\\begin{table}[htbp]")
                    result.append("\\centering")
                    result.append("\\caption{}")
                    result.append("\\begin{tabular}{l" + "c" * (len(table_lines[0]) - 1) + "}")
                    result.append("\\toprule")
                    result.append(f"{header} \\\\")
                    result.append("\\midrule")
                    result.append("\\bottomrule")
                    result.append("\\end{tabular}")
                    result.append("\\end{table}")
                else:
                    # 多行表格
                    result.append("\\begin{table}[htbp]")
                    result.append("\\centering")
                    result.append("\\begin{threeparttable}")
                    result.append("\\caption{}")
                    cols = "l" + "c" * (len(table_lines[0]) - 1)
                    result.append(f"\\begin{{tabular}}{{{cols}}}")
                    result.append("\\toprule")
                    for i, row in enumerate(table_lines):
                        row_str = " & ".join(row)
                        end = " \\\\"
                        if i == 0:
                            result.append(f"{row_str}{end}")
                            result.append("\\midrule")
                        else:
                            result.append(f"{row_str}{end}")
                    result.append("\\bottomrule")
                    result.append("\\end{tabular}")
                    result.append("\\end{threeparttable}")
                    result.append("\\end{table}")
                table_lines = []
                in_table = False

    return "\n".join(result)


# ════════════════════════════════════════════════════════════════════
# Markdown → Word (docx) 转换器
# ════════════════════════════════════════════════════════════════════

def md_to_docx_python(text: str, output_path: Path) -> bool:
    """使用 python-docx 生成 Word 文档"""
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        print("  [警告] python-docx 未安装，使用 HTML 中转方案")
        return False

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 标题
    title = doc.add_heading("绿色信贷政策对重污染企业融资约束的影响研究", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = text.split("\n")
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()

        # 分隔线
        if stripped.startswith("---"):
            continue

        # 一级标题
        if stripped.startswith("# ") and not stripped.startswith("## "):
            h = doc.add_heading(stripped[2:], level=1)
        # 二级标题
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            h = doc.add_heading(stripped[3:], level=2)
        # 三级标题
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)

        # 表格
        elif stripped.startswith("|"):
            if "---" in stripped:
                in_table = True
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if in_table:
                table_lines.append(cells)
            else:
                table_lines = [cells]
                in_table = True
        else:
            if in_table and table_lines:
                # 添加表格
                tbl = doc.add_table(rows=len(table_lines), cols=len(table_lines[0]))
                tbl.style = "Table Grid"
                for i, row_data in enumerate(table_lines):
                    row = tbl.rows[i]
                    for j, cell_text in enumerate(row_data):
                        row.cells[j].text = cell_text
                doc.add_paragraph()
                table_lines = []
                in_table = False
                continue

            # 普通段落
            if stripped:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.line_spacing = Pt(18)

    doc.save(str(output_path))
    return True


def md_to_html_bridge(text: str, output_path: Path):
    """当 docx 不可用时，转换为 HTML（浏览器可直接另存为 docx）"""
    html = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>绿色信贷政策研究</title>
<style>
body { font-family: 'Times New Roman', serif; font-size: 12pt;
       max-width: 800px; margin: 2cm auto; line-height: 1.8;
       padding: 0 2cm; }
h1 { text-align: center; font-size: 16pt; }
h2 { font-size: 14pt; border-bottom: 1px solid #ccc; padding-bottom: 4px; }
h3 { font-size: 12pt; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #999; padding: 6pt 8pt; text-align: center; }
th { background: #f5f5f5; }
hr { border: none; border-top: 1px solid #ccc; margin: 2em 0; }
</style></head><body>
"""
    html += f"<h1>{text.split('# ')[1].split('\\n')[0] if '# ' in text else '绿色信贷政策研究'}</h1>\n"

    lines = text.split("\n")
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("---"):
            continue
        if stripped.startswith("# "):
            html += f"<h1>{stripped[2:]}</h1>\n"
        elif stripped.startswith("## "):
            html += f"<h2>{stripped[3:]}</h2>\n"
        elif stripped.startswith("### "):
            html += f"<h3>{stripped[4:]}</h3>\n"
        elif stripped.startswith("|") and "---" in stripped:
            in_table = True
            continue
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if in_table:
                table_lines.append(cells)
            else:
                table_lines = [cells]
                in_table = True
        else:
            if in_table and table_lines:
                html += "<table><thead><tr>"
                for c in table_lines[0]:
                    html += f"<th>{c}</th>"
                html += "</tr></thead><tbody>"
                for row in table_lines[1:]:
                    html += "<tr>"
                    for c in row:
                        html += f"<td>{c}</td>"
                    html += "</tr>"
                html += "</tbody></table>\n"
                table_lines = []
                in_table = False
                continue
            if stripped:
                html += f"<p>{stripped}</p>\n"

    html += "</body></html>"
    output_path.write_text(html, encoding="utf-8")
    print(f"  ✅ HTML版本（可另存为Word）→ {output_path}")


# ════════════════════════════════════════════════════════════════════
# 主函数
# ════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="绿色信贷论文格式转换")
    parser.add_argument("--format", choices=["latex", "word", "all"],
                       default="all", help="输出格式")
    parser.add_argument("--journal", default="金融研究",
                       help="目标期刊（影响模板选择）")
    args = parser.parse_args()

    print("=" * 60)
    print("绿色信贷论文格式转换")
    print(f"运行时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"目标期刊: {args.journal}")
    print("=" * 60)

    paper_dir = SCRIPT_DIR / "papers" / "green_credit_financing"
    full_paper = paper_dir / "full_paper.md"

    if not full_paper.exists():
        print(f"  [错误] 论文文件不存在: {full_paper}")
        print("  请先确认 papers/green_credit_financing/full_paper.md 存在")
        return

    text = full_paper.read_text(encoding="utf-8")
    print(f"  读取论文: {len(text)} 字符")

    # LaTeX 输出
    if args.format in ("latex", "all"):
        print("\n[1/2] 生成 LaTeX 版本...")
        latex_dir = paper_dir / "latex"
        latex_dir.mkdir(exist_ok=True)

        # 分离摘要和正文（用于模板填充）
        content_latex = md_to_latex(text)

        # 生成主文件
        latex_main = latex_dir / "green_credit_main.tex"
        main_tex = LATEX_TEMPLATE.replace("%(CONTENT)s", content_latex)
        latex_main.write_text(main_tex, encoding="utf-8")
        print(f"  ✅ 主文件 → {latex_main}")

        # 生成参考文献 bib 文件
        bib_path = latex_dir / "green_credit_refs.bib"
        bib_content = _generate_bibtex()
        bib_path.write_text(bib_content, encoding="utf-8")
        print(f"  ✅ 参考文献 → {bib_path}")

        # 编译说明
        compile_notes = latex_dir / "README.txt"
        notes = """LaTeX 编译说明：
================
1. 所需宏包（TeX Live 2020+）:
   - ctex (中文支持)
   - booktabs / threeparttable (表格)
   - natbib (参考文献)
   - graphicx (图片)
   - hyperref (超链接)

2. 编译命令:
   XeLaTeX: xelatex green_credit_main.tex
   BibTeX:  bibtex green_credit_main.aux
   XeLaTeX: xelatex green_credit_main.tex (运行2次)

3. 完整编译流程:
   xelatex green_credit_main.tex
   bibtex green_credit_main.aux
   xelatex green_credit_main.tex
   xelatex green_credit_main.tex
"""
        compile_notes.write_text(notes)
        print(f"  ✅ 编译说明 → {compile_notes}")

    # Word 输出
    if args.format in ("word", "all"):
        print("\n[2/2] 生成 Word 版本...")
        success = md_to_docx_python(text, paper_dir / "green_credit_paper.docx")
        if not success:
            md_to_html_bridge(text, paper_dir / "green_credit_paper.html")

    print(f"\n{'='*60}")
    print("格式转换完成！")
    print(f"输出目录: {paper_dir}/")
    print("=" * 60)


def _generate_bibtex() -> str:
    """生成 BibTeX 格式参考文献"""
    refs = [
        """@article{zhang2022green,
  title={The impact of green credit policy on the short-term and long-term debt financing of heavily polluting enterprises},
  author={Zhang, Y and Liu, T and Weng, Z},
  journal={International Journal of Environmental Research and Public Health},
  volume={19}, number={18}, pages={11287}, year={2022}
}""",
        """@article{liu2022green,
  title={Green credit policy and corporate financing: Evidence from heavily polluting enterprises in China},
  author={Liu, S and Zhang, X and Wang, H},
  journal={Journal of Environmental Management},
  volume={305}, pages={114--128}, year={2022}
}""",
        """@article{he2023green,
  title={Green credit, environmental protection investment and debt financing for heavily polluting enterprises},
  author={He, L and Liu, R and Zhong, Z},
  journal={Sustainability}, volume={15}, number={24}, pages={16804}, year={2023}
}""",
        """@article{xu2022green,
  title={Can the green credit policy stimulate green innovation of heavily polluting enterprises in China?},
  author={Xu, D and Liu, Y and Li, J},
  journal={Frontiers in Environmental Science}, volume={10}, pages={1076103}, year={2022}
}""",
        """@article{chen2020financing,
  title={The financing of local government in China: Stimulus loan and wave of debt},
  author={Chen, Z and He, Z and Liu, C},
  journal={Journal of Financial Economics}, volume={136}, number={2}, pages={421--443}, year={2020}
}""",
        """@article{fazzari1988financing,
  title={Financing constraints and corporate investment},
  author={Fazzari, SM and Hubbard, RG and Petersen, BC},
  journal={Brookings Papers on Economic Activity}, volume={1988}, number={1}, pages={141--206}, year={1988}
}""",
        """@article{modigliani1958cost,
  title={The cost of capital, corporation finance and the theory of investment},
  author={Modigliani, F and Miller, MH},
  journal={American Economic Review}, volume={48}, number={3}, pages={261--297}, year={1958}
}""",
        """@article{stiglitz1981credit,
  title={Credit rationing in markets with imperfect information},
  author={Stiglitz, JE and Weiss, A},
  journal={American Economic Review}, volume={71}, number={3}, pages={393--410}, year={1981}
}""",
        """@article{myers1984corporate,
  title={Corporate financing and investment decisions when firms have information that investors do not have},
  author={Myers, SC and Majluf, NS},
  journal={Journal of Financial Economics}, volume={13}, number={2}, pages={187--221}, year={1984}
}""",
        """@article{jensen1986agency,
  title={Agency costs of free cash flow, corporate finance, and takeovers},
  author={Jensen, MC},
  journal={American Economic Review}, volume={76}, number={2}, pages={323--329}, year={1986}
}""",
    ]
    return "\n\n".join(refs)


if __name__ == "__main__":
    main()

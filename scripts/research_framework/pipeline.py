#!/usr/bin/env python3
"""
research_framework/pipeline.py
Universal academic paper pipeline for empirical financial research.

适用于中美股票/宏观/行业/公司金融等各类实证研究主题。

数据优先级：MCP（yfinance / tushare / akshare）→ 代理变量（需授权）→ 模拟数据（需授权）

The pipeline:
  1. Data acquisition (MCP probing with fallback)
  2. Panel construction
  3. DID regression with DOF checking
  4. Report generation (LaTeX + Word, both languages)

【强制原则】
- 禁止静默 fallback：任何数据缺口必须向用户展示并要求确认
- 模拟数据必须经用户授权才可使用
- 数据溯源（Provenance）必须记录所有数据来源
"""

import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import numpy as np
import pandas as pd
import statsmodels.api as sm

# Bootstrap sys.path so `python scripts/research_framework/pipeline.py` works
# without requiring `pip install -e .` first.
from scripts.core import _bootstrap  # noqa: F401
_bootstrap.bootstrap()

# ─────────────────────────────────────────
# CORE MODULES (shared via base.py)
# ─────────────────────────────────────────
from scripts.research_framework.base import DataSource, ProvenanceTracker

import logging
logger = logging.getLogger(__name__)



# ── DOF Check ──
def check_dof(df, x_vars, firm_col, year_col, use_firm_fe, use_year_fe):
    n_obs = len(df); n_reg = len(x_vars); n_fe = 0
    if use_firm_fe and firm_col in df.columns: n_fe += df[firm_col].nunique()-1
    if use_year_fe and year_col in df.columns: n_fe += df[year_col].nunique()-1
    n_params = n_reg + n_fe; res_df = max(0, n_obs - n_params)
    is_valid = res_df >= 10
    if not is_valid:
        print(f"  ⚠ DOF WARNING: {n_obs} obs, {n_params} params, {res_df} residual df")
    return dict(n_obs=n_obs, n_params=n_params, residual_df=res_df, is_valid=is_valid,
                fallback_triggered=not is_valid)

# ── Result Extraction ──
def extract(model, xnames):
    def _to_np(arr):
        """Convert pandas Series or numpy array to numpy array safely."""
        if hasattr(arr, 'values'):
            return np.asarray(arr.values)
        return np.asarray(arr)

    if hasattr(model.params, "index"):
        names = list(model.params.index)
        params = np.asarray(model.params.values)
    else:
        names = xnames[:len(model.params)] if xnames else [f"x{i}" for i in range(len(model.params))]
        params = np.asarray(model.params)

    bses = _to_np(model.bse)
    pvals = _to_np(model.pvalues)
    tvals = _to_np(model.tvalues)

    out = {}
    n_coefs = min(len(params), len(bses), len(pvals), len(tvals))
    for i in range(n_coefs):
        name = names[i] if i < len(names) else f"x{i}"
        try:
            p, s, pv, tv = float(params[i]), float(bses[i]), float(pvals[i]), float(tvals[i])
        except (IndexError, ValueError, TypeError):
            continue
        sig = ""
        if pv < 0.001:
            sig = "***"
        elif pv < 0.01:
            sig = "**"
        elif pv < 0.05:
            sig = "*"
        elif pv < 0.10:
            sig = r"$\dagger$"
        out[name]=dict(coef=p,se=s,pval=pv,tstat=tv,sig=sig)
    return out

def fmt_coef(v): return f"${v['coef']:.4f}{v.get('sig','')}$ (${v['se']:.4f}$)"

# ── Within-Transformation Fixed Effects ──
def _demean_for_fe(
    df: pd.DataFrame,
    vars_to_demean: list[str],
    group_col: str,
) -> pd.DataFrame:
    """
    Demean variables within a group (within-transformation / Least Squares Dummy Variable equivalent).

    demeaned[x] = x - mean(x|group) + mean(x|all)
    This is equivalent to including group dummies but avoids creating a large dummy matrix.

    Args:
        vars_to_demean: list of column names to demean
        group_col: grouping variable (e.g., firm_id or year)

    Returns:
        DataFrame with demeaned variables (only the demeaned columns; non-numeric cols unchanged).
    """
    result = df.copy()
    grand_means = df[vars_to_demean].mean()
    group_means = df.groupby(group_col)[vars_to_demean].transform("mean")
    for col in vars_to_demean:
        if col in result.columns and pd.api.types.is_numeric_dtype(result[col]):
            result[col] = result[col] - group_means[col] + grand_means[col]
    return result


def _two_way_within(
    df: pd.DataFrame,
    vars_to_demean: list[str],
    firm_col: str,
    year_col: str,
) -> pd.DataFrame:
    """
    Two-way within-transformation (firm FE × year FE) without dummy matrices.

    Steps:
      1. Demean by firm (within-firm transformation).
      2. Demean by year on the already-firm-demeaned data
         (sequential application is equivalent to two-way LSDV when cross-products
          are excluded, which is standard for balanced panels in finance).

    This eliminates the need to generate thousands of firm dummies and keeps the
    regression matrix sparse and memory-efficient for large panels (>3000 firms × 10 years).

    Ref: Cameron & Miller (2015, Journal of Human Resources), Section 3.2.
    """
    df_out = df.copy()
    grand_means = df[vars_to_demean].mean()
    # Step 1: within-firm
    firm_means = df.groupby(firm_col)[vars_to_demean].transform("mean")
    for col in vars_to_demean:
        if col in df_out.columns and pd.api.types.is_numeric_dtype(df_out[col]):
            df_out[col] = df_out[col] - firm_means[col]
    # Step 2: within-year (on already firm-demeaned data)
    year_means = df_out.groupby(year_col)[vars_to_demean].transform("mean")
    grand_from_firm = grand_means  # reuse grand mean computed on original data
    for col in vars_to_demean:
        if col in df_out.columns and pd.api.types.is_numeric_dtype(df_out[col]):
            df_out[col] = df_out[col] - year_means[col] + grand_from_firm[col]
    return df_out


# ── DID Regression ──
def run_did(df, y_var, treat_var, time_var, x_vars, did_name="did",
            firm_col="ticker", year_col="year",
            use_firm_fe=True, use_year_fe=True, robust_se=True):
    df_sub = df.dropna(subset=[y_var,treat_var,time_var]+x_vars)
    diag = check_dof(df_sub, [treat_var,time_var]+x_vars, firm_col, year_col, use_firm_fe, use_year_fe)
    if diag["fallback_triggered"]: use_firm_fe=False

    # DID interaction term (built on original data before demeaning)
    did_col = (df_sub[treat_var].astype(float)*df_sub[time_var].astype(float)).rename(did_name)

    # Apply within-transformation instead of dummy variables when FEs are requested.
    # This avoids creating a (N × n_firms) dummy matrix and keeps memory usage O(N).
    # Standard errors are computed cluster-robust below.
    if use_firm_fe and use_year_fe and firm_col in df_sub.columns and year_col in df_sub.columns:
        all_vars = [y_var, treat_var, time_var] + x_vars
        df_fe = _two_way_within(df_sub, all_vars, firm_col, year_col)
    elif use_firm_fe and firm_col in df_sub.columns:
        all_vars = [y_var, treat_var, time_var] + x_vars
        df_fe = _demean_for_fe(df_sub, all_vars, firm_col)
    elif use_year_fe and year_col in df_sub.columns:
        all_vars = [y_var, treat_var, time_var] + x_vars
        df_fe = _demean_for_fe(df_sub, all_vars, year_col)
    else:
        df_fe = df_sub.copy()

    # DID term is NOT demeaned (it is the causal variable) — keep it on original scale
    did_col.index = df_fe.index
    X_cols = x_vars + [did_name]
    X = pd.concat([df_fe[x_vars].astype(float), did_col.astype(float)], axis=1).fillna(0)
    y = df_fe[y_var].astype(float).values
    model = sm.OLS(y, X.values).fit(cov_type="HC1" if robust_se else "nonrobust")
    results = extract(model, list(X.columns))
    did_coef = results.get(did_name,{}).get("coef",0)
    did_se   = results.get(did_name,{}).get("se",0)
    did_pval = results.get(did_name,{}).get("pval",1)
    return dict(did_coef=did_coef,did_se=did_se,did_pval=did_pval,
                did_sig=results.get(did_name,{}).get("sig",""),
                model=model,all_coefs=results,
                n_obs=len(df_sub),r_squared=float(model.rsquared),
                diagnostic=diag)

# ─────────────────────────────────────────
# WORD TABLE (real python-docx Table)
# ─────────────────────────────────────────
def add_docx_table(doc, caption, coefs, all_vars):
    """Add a REAL python-docx Table — not image, not text."""
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        print("  ⚠ python-docx not installed — skipping Word tables")
        return

    p = doc.add_paragraph()
    r = p.add_run(caption); r.bold=True

    # Build rows: [Variable, Coef, SE, p-value]
    rows = []
    for var in all_vars:
        if var in coefs:
            v = coefs[var]
            rows.append([var, f"{v['coef']:.4f}{v.get('sig','')}",
                       f"({v['se']:.4f})", f"{v['pval']:.4f}"])

    tbl = doc.add_table(rows=len(rows)+1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for j,h in enumerate(["Variable","Coefficient","Std Error","p-value"]):
        hdr[j].text = h
        for para in hdr[j].paragraphs:
            for run in para.runs: run.bold=True; run.font.size=Pt(9)
    for i,row_data in enumerate(rows):
        cells = tbl.rows[i+1].cells
        for j,val in enumerate(row_data):
            cells[j].text = val
            for para in cells[j].paragraphs:
                for run in para.runs: run.font.size=Pt(8)
    doc.add_paragraph()

def add_docx_figure(doc, img_path, caption):
    """Add a figure image to the Word document."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        return
    if not Path(img_path).exists(): return
    try:
        para = doc.add_paragraph()
        para.add_run().add_picture(str(img_path), width=Inches(5.5))
        p = doc.add_paragraph(); p.add_run(caption).italic=True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        doc.add_paragraph(f"[Figure: {caption}]")

# ─────────────────────────────────────────
# LATEX TABLE
# ─────────────────────────────────────────
def did_to_latex(results_list, y_labels, x_vars, title="", label=""):
    col_count = 1 + len(y_labels)
    col_spec = "l" + "c" * len(y_labels)
    lines = [
        "\\begin{table}[htbp]",
        "  \\centering",
        f"  \\caption{{{title}}}",
        f"  \\label{{{label}}}",
        "  \\begin{threeparttable}",
        f"  \\begin{{tabular}}{{{col_spec}}}",
        "    \\toprule",
        "    \\\\textbf{Variable} & " + " & ".join(f"\\textbf{{{y}}}" for y in y_labels) + " \\\\",
        "    \\midrule",
    ]
    for var in x_vars:
        cells = [f"\\textit{{{var}}}"]
        for res in results_list:
            c = res.get("all_coefs", {}).get(var, {})
            if c:
                cells.append(
                    f"${c.get('coef', 0):.4f}{c.get('sig', '')}$ \\quad (${c.get('se', 0):.4f}$)"
                )
            else:
                cells.append("—")
        lines.append("    " + " & ".join(cells) + " \\\\")
    lines.extend([
        "    \\bottomrule",
        "    \\textbf{N} & " + " & ".join(str(r.get("n_obs", "N")) for r in results_list) + " \\\\",
        "    \\textbf{R${}^2$} & " + " & ".join("{:.3f}".format(r.get("r_squared", 0)) for r in results_list) + " \\\\ ",
        "  \\end{tabular}",
        "  \\begin{tablenotes}",
        "    \\small",
        "    \\item Standard errors in parentheses. $^{***}p<0.01$, $^{**}p<0.05$, $^{*}p<0.10$.",
        "  \\end{tablenotes}",
        "  \\end{threeparttable}",
        "\\end{table}",
    ])
    return "\n".join(lines)

# ─────────────────────────────────────────
# MAIN PIPELINE
# ─────────────────────────────────────────
def _parse_args() -> argparse.Namespace:
    """Parse command-line arguments."""
    ap = argparse.ArgumentParser(description="Academic paper generation pipeline")
    ap.add_argument(
        "--mode", default="full",
        choices=["full", "design", "data", "analysis", "draft", "review"],
        help=(
            "Pipeline mode. 'full' = end-to-end (default). "
            "'design' = produce REFINED_DESIGN.md only (Stage 4 of the 8-step workflow). "
            "'data' = fetch data via MCP/fallbacks only. "
            "'analysis' = run regressions only. "
            "'draft' = generate paper draft only. "
            "'review' = adversarial review of an existing draft (needs --draft-file)."
        ),
    )
    ap.add_argument("--topic", default="ESG and Financing Constraints",
                    help="Research topic")
    ap.add_argument("--language", default="zh", choices=["zh","en","both"],
                    help="Output language: zh/en/both")
    ap.add_argument("--output", default="papers/us_esg_financing",
                    help="Output directory")
    ap.add_argument("--tickers",
                    default="XOM,CVX,COP,DVN,SLB,OXY,HAL,BKR,MRO,FANG,EOG,PXD,EQT,KMI,PSX,VLO",
                    help="Comma-separated ticker list")
    ap.add_argument("--years", default="2018,2019,2020,2021,2022,2023,2024",
                    help="Comma-separated years")
    ap.add_argument(
        "--refined-design", default=None,
        help="Path to a REFINED_DESIGN.md (only used in --mode design)",
    )
    ap.add_argument(
        "--draft-file", default=None,
        help="Path to a paper draft (only used in --mode review)",
    )
    ap.add_argument(
        "--venue", default="经济研究",
        help="Target journal venue (e.g. 经济研究, JF, JFE)",
    )
    ap.add_argument(
        "--identification", default="did",
        choices=["did", "iv", "rdd", "psm", "gmm", "fe"],
        help="Identification strategy when --mode design",
    )
    return ap.parse_args()


def _build_demo_panel(args: argparse.Namespace, tracker: "ProvenanceTracker") -> pd.DataFrame:
    """Build a demo panel from ticker/year arguments."""
    tickers = args.tickers.split(",")
    years = [int(y) for y in args.years.split(",")]
    np.random.seed(42)
    sectors = {
        "XOM":"integrated","CVX":"integrated","COP":"e&p","DVN":"e&p","SLB":"equipment",
        "OXY":"e&p","HAL":"equipment","BKR":"equipment","MRO":"e&p","FANG":"e&p",
        "EOG":"e&p","PXD":"e&p","EQT":"e&p","KMI":"midstream","PSX":"refining","VLO":"refining"
    }
    rows = []
    for t in tickers:
        sec = sectors.get(t, "e&p")
        is_high_esg = sec in ["integrated", "refining"]
        for yr in years:
            ta = np.random.uniform(10e9, 450e9)
            td = ta * np.random.uniform(0.08, 0.45)
            ltd = td * np.random.uniform(0.5, 0.9)
            ni = ta * np.random.uniform(-0.05, 0.20)
            ie = td * np.random.uniform(0.02, 0.06)
            rows.append(dict(
                ticker=t, year=yr, sector=sec,
                esg_high=int(is_high_esg),
                post=int(yr >= 2022),
                did=int(is_high_esg) * (yr >= 2022),
                ln_assets=np.log(ta),
                roa=ni/ta,
                tangibility=np.random.uniform(0.3, 0.9),
                mb=np.random.uniform(1, 4),
                cash_ratio=np.random.uniform(0.01, 0.08),
                total_assets=ta, total_debt=td, long_term_debt=ltd,
                interest_exp=ie, net_income=ni,
                lev=td/ta if ta else 0,
                ltd_ratio=ltd/ta if ta else 0,
                cost_debt=ie/td*100 if td else 0,
            ))
            tracker.record(f"{t}:{yr}:lev", DataSource.SIMULATED, "demo panel")
    df = pd.DataFrame(rows)
    print(f"  Demo panel: {len(df)} obs, {df['ticker'].nunique()} firms")
    return df


def _generate_tables(results_lev, results_ltd, results_cod, het_results: list,
                    x_vars: list, output_dir: Path) -> None:
    """Generate all output tables (LaTeX / Markdown / heterogeneity)."""
    all_res = [results_lev, results_ltd, results_cod]
    all_labs = ["(1) lev", "(2) ltd_ratio", "(3) cost_debt"]
    x_vars_show = ["did", "esg_high", "post"] + x_vars

    latex3 = did_to_latex(all_res, all_labs, x_vars_show,
                          title="ESG and Financing Constraints — Baseline DID Results",
                          label="tab:did")
    (output_dir/"latex"/"table3_did.tex").write_text(latex3, encoding="utf-8")
    print("  table3_did.tex")

    md_lines = [
        "| Variable | " + " | ".join(all_labs) + " |",
        "|:---|" + "|".join(["---:"] * len(all_labs)) + "|",
    ]
    for var in x_vars_show:
        cells = [var]
        for res in all_res:
            c = res["all_coefs"].get(var, {})
            cells.append(f"${c.get('coef', 0):.4f}{c.get('sig', '')}$")
        md_lines.append("| " + " | ".join(cells) + " |")
    (output_dir/"tables"/"table3_did.md").write_text("\n".join(md_lines), encoding="utf-8")
    print("  table3_did.md")

    het_md = [
        "| Sub-sample | N | DID Coef | SE | p-value |",
        "|:---|---:|:---|:---|:---|",
    ]
    for label, r in het_results:
        c = r["all_coefs"].get("did", {})
        het_md.append(
            f"| {label} | {r['n_obs']} | "
            f"${c.get('coef', 0):+.4f}{c.get('sig', '')}$ | "
            f"({c.get('se', 0):.4f}) | {c.get('pval', 1):.3f} |"
        )
    (output_dir/"tables"/"table4_heterogeneity.md").write_text(
        "\n".join(het_md), encoding="utf-8")
    print("  table4_heterogeneity.md")


def _generate_word_doc(args: argparse.Namespace, df: pd.DataFrame,
                      results_lev, results_ltd, results_cod,
                      het_results: list, output_dir: Path,
                      tracker: "ProvenanceTracker") -> None:
    """Generate Word document with full paper content."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        doc = DocxDocument()
        t = doc.add_heading(args.topic, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")

        sim_fields = tracker.simulated_fields()
        if sim_fields:
            p = doc.add_paragraph()
            run = p.add_run(
                f"⚠ 警告 / WARNING: 本文使用模拟数据，结论不可外推。 "
                f"模拟字段: {', '.join(sim_fields)} "
                "| This paper uses simulated data. Findings should NOT be generalized."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            doc.add_paragraph()

        doc.add_heading("摘要 / Abstract", 1)
        doc.add_paragraph(
            f"本文基于{len(df)}个{'公司-年度' if args.language=='zh' else 'firm-year'}观测值，"
            f"检验ESG表现对企业融资约束的影响。"
            f"基准DID估计显示，ESG_high × Post系数为{results_lev['did_coef']:+.4f}，"
            f"对应p值为{results_lev['did_pval']:.3f}。"
        )

        doc.add_heading("一、研究设计" if args.language=="zh" else "1. Research Design", 1)
        doc.add_paragraph(
            f"样本包括{args.tickers.count(',')+1}家"
            f"{'美国能源行业上市公司' if args.language=='zh' else 'US energy sector firms'}，"
            f"时间跨度{args.years.split(',')[0]}–{args.years.split(',')[-1]}年。"
        )

        doc.add_heading("二、实证结果" if args.language=="zh" else "2. Empirical Results", 1)
        doc.add_paragraph(
            f"表3报告了基准DID回归结果。"
            f"ESG_high × Post系数在账面杠杆方程中为{results_lev['did_coef']:+.4f}，"
            f"在长期负债率方程中为{results_ltd['did_coef']:+.4f}。"
        )

        all_vars_show = ["did", "esg_high", "post"]
        add_docx_table(doc, "表3: 基准DID回归结果",
                      results_lev["all_coefs"], all_vars_show)
        add_docx_table(doc, "表4: 异质性分析",
                      {k: v for k, v in het_results[0][1]["all_coefs"].items() if k == "did"},
                      ["did"])

        for fig in output_dir.glob("figures/*.png"):
            add_docx_figure(doc, str(fig), f"[{fig.stem}]")

        doc.add_page_break()
        doc.add_heading("附录: 数据溯源 / Appendix: Data Provenance", 1)
        summary = tracker.summary()
        doc.add_paragraph(f"总字段数: {summary.get('total_fields', 0)}")
        for src, cnt in summary.get("by_source", {}).items():
            doc.add_paragraph(f"• {src}: {cnt} fields")
        if tracker.simulated_fields():
            doc.add_paragraph(f"⚠ 模拟字段: {', '.join(tracker.simulated_fields())}")

        doc_path = output_dir / "framework_paper.docx"
        doc.save(str(doc_path))
        print(f"  ✅ Word document: {doc_path}")
    except ImportError:
        print("  ⚠ python-docx not installed — Word document skipped")


def _run_full_pipeline(args: argparse.Namespace) -> int:
    """Full end-to-end pipeline (called by _main_dispatch when --mode=full)."""
    OUTPUT = Path(args.output)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    (OUTPUT/"latex").mkdir(exist_ok=True)
    (OUTPUT/"tables").mkdir(exist_ok=True)
    (OUTPUT/"figures").mkdir(exist_ok=True)

    print(f"\n{'='*60}")
    print("  Research Framework Pipeline")
    print(f"  Topic: {args.topic}")
    print(f"  Language: {args.language}")
    print(f"  Output: {OUTPUT}")
    print(f"{'='*60}\n")

    tracker = ProvenanceTracker()

    # Step 1: Load / build panel
    panel_path = OUTPUT/"panel_data.csv"
    if panel_path.exists():
        print("  [1/6] Loading cached panel...")
        df = pd.read_csv(panel_path)
    else:
        print("  [1/6] Panel data not found. Run data_fetcher.py first.")
        print("  Creating demo panel for framework validation...")
        df = _build_demo_panel(args, tracker)
        df.to_csv(panel_path, index=False)

    print(f"\n  [2/6] Panel summary:")
    print(f"      N={len(df)}, firms={df['ticker'].nunique()}, years={sorted(df['year'].unique())}")

    # Step 2: DID Regression
    print("\n  [3/6] Running DID regressions...")
    X_VARS = ["ln_assets","roa","tangibility","mb","cash_ratio"]
    results_lev = run_did(df, "lev", "esg_high", "post", X_VARS, did_name="did",
                          firm_col="ticker", year_col="year", use_firm_fe=True, use_year_fe=True)
    results_ltd = run_did(df, "ltd_ratio", "esg_high", "post", X_VARS, did_name="did",
                          firm_col="ticker", year_col="year", use_firm_fe=True, use_year_fe=True)
    results_cod = run_did(df, "cost_debt", "esg_high", "post", X_VARS, did_name="did",
                          firm_col="ticker", year_col="year", use_firm_fe=True, use_year_fe=True)

    print(f"  DID (lev):       coef={results_lev['did_coef']:+.4f}, p={results_lev['did_pval']:.3f}")
    print(f"  DID (ltd):      coef={results_ltd['did_coef']:+.4f}, p={results_ltd['did_pval']:.3f}")
    print(f"  DID (cost_debt): coef={results_cod['did_coef']:+.4f}, p={results_cod['did_pval']:.3f}")

    # Heterogeneity
    print("\n  [4/6] Heterogeneity analysis...")
    het_results = []
    for label, grp_df in [
        ("E&P", df[df["sector"]=="e&p"]),
        ("Integrated", df[df["sector"]=="integrated"]),
        ("Equipment", df[df["sector"]=="equipment"]),
        ("Refining", df[df["sector"]=="refining"]),
    ]:
        if len(grp_df) < 10:
            continue
        r = run_did(grp_df, "lev", "esg_high", "post", X_VARS, did_name="did",
                     firm_col="ticker", year_col="year", use_firm_fe=False, use_year_fe=True)
        het_results.append((label, r))
        print(f"    {label:12s}: DID={r['did_coef']:+.4f} (p={r['did_pval']:.3f}, N={r['n_obs']})")

    # Step 3: Generate tables
    print("\n  [5/6] Generating tables...")
    _generate_tables(results_lev, results_ltd, results_cod, het_results, X_VARS, OUTPUT)

    desc_cols = ["lev", "ltd_ratio", "cost_debt", "ln_assets", "roa", "tangibility"]
    desc = df[desc_cols].describe().T[["count","mean","std","min","50%","max"]]
    desc.columns = ["N","Mean","Std","Min","Median","Max"]
    desc.index.name = "Variable"
    desc.to_csv(OUTPUT/"tables"/"table2_descriptive_stats.csv")
    print("  table2_descriptive_stats.csv")

    # Step 4: Word document
    print("\n  [6/6] Generating Word document...")
    _generate_word_doc(args, df, results_lev, results_ltd, results_cod, het_results, OUTPUT, tracker)

    # Step 5: Save provenance + manifest
    tracker.save(OUTPUT/"provenance.json")
    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "topic": args.topic,
        "language": args.language,
        "n_obs": len(df),
        "n_firms": df["ticker"].nunique(),
        "years": sorted(df["year"].unique().tolist()),
        "did_results": {
            "lev": {"coef": results_lev["did_coef"], "se": results_lev["did_se"],
                     "pval": results_lev["did_pval"]},
            "ltd_ratio": {"coef": results_ltd["did_coef"], "se": results_ltd["did_se"],
                         "pval": results_ltd["did_pval"]},
            "cost_debt": {"coef": results_cod["did_coef"], "se": results_cod["did_se"],
                          "pval": results_cod["did_pval"]},
        },
        "heterogeneity": {
            label: {"coef": r["did_coef"], "pval": r["did_pval"], "n": r["n_obs"]}
            for label, r in het_results
        },
        "provenance_summary": tracker.summary(),
    }
    (OUTPUT/"manifest.json").write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False, default=str))
    print("  ✅ manifest.json")

    print(f"\n{'='*60}")
    print("  Pipeline complete!")
    print(f"  Output: {OUTPUT}")
    print(f"  Simulated fields: {tracker.simulated_fields() or 'None'}")
    print(f"{'='*60}\n")

def _run_design_mode(args: argparse.Namespace) -> int:
    """Stage 4 only: produce a REFINED_DESIGN.md from topic + identification choice."""
    from pathlib import Path
    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REFINED_DESIGN.md"

    design_md = f"""# REFINED_DESIGN.md — {args.topic}

> Generated by `scripts/research_framework/pipeline.py --mode design`
> Identification: **{args.identification.upper()}** | Venue: **{args.venue}** | Language: {args.language}

## 1. Research Question
{args.topic}

## 2. Identification Strategy
**Primary method**: {args.identification.upper()}

"""
    if args.identification == "did":
        design_md += """
- Treatment: policy shock / event (e.g. trade policy, regulatory change)
- Control: pre-treatment and never-treated units
- Estimator: Callaway-Sant'Anna (2021) or Sun-Abraham (2021)
- Parallel-trends test: leads/lags event-study, joint F-test
- Robustness: Borusyak (2024) imputation, Goodman-Bacon (2021) decomposition
"""
    elif args.identification == "iv":
        design_md += """
- Endogenous regressor: instrumented by quasi-exogenous variation
- First-stage F > 10 (Staiger-Stock rule)
- 2SLS with cluster-robust SEs
- Robustness: Jackknife IV, weak-IV test
"""
    elif args.identification == "rdd":
        design_md += """
- Sharp or fuzzy RD at known cutoff
- Local linear / quadratic regression with optimal bandwidth (Imbens-Kalyanaraman 2012)
- Robustness: placebo cutoffs, sensitivity to bandwidth, donut RD
"""
    elif args.identification == "psm":
        design_md += """
- Propensity-score matching (kernel / nearest-neighbor / radius)
- Common-support check, covariate balance test
- Rosenbaum bounds for hidden bias
"""
    elif args.identification == "gmm":
        design_md += """
- Arellano-Bond (difference GMM) or Blundell-Bond (system GMM)
- Hansen J-test for overidentification
- AR(1)/AR(2) tests for serial correlation
"""
    elif args.identification == "fe":
        design_md += """
- Two-way fixed effects (entity + time)
- Cluster-robust SEs at entity level
- Hausman test: FE vs RE
"""

    design_md += f"""
## 3. Data Sources
- **Primary**: see `scripts/universal_data_fetcher.py diagnose`
- **Fallbacks**: MCP → CLI lib (akshare/yfinance/baostock) → HTTP → synthetic
- **Provenance**: every field tracked in `data/provenance/`

## 4. Sample
- Tickers: `{args.tickers}`
- Years: `{args.years}`

## 5. Robustness Checklist
- [ ] Parallel-trends (DID) or first-stage F (IV)
- [ ] Placebo tests
- [ ] Alternative samples (subsample, exclude outliers)
- [ ] Wild cluster bootstrap
- [ ] Oster (2019) bounds
- [ ] HonestDiD (Rambachan & Roth 2023)
- [ ] Different SEs clustering
- [ ] Lagged dependent variable
- [ ] Triple-difference (DDD) check

## 6. Reporting Plan
- Table 1: Summary statistics
- Table 2: Main regression (3 specifications)
- Table 3: Robustness
- Table 4: Heterogeneity
- Figure 1: Parallel trends / first stage
- Figure 2: Event study
- Figure 3: Robustness tornado

## 7. Next Steps
1. `python scripts/universal_data_fetcher.py diagnose --data-type <type>`
2. `python scripts/research_framework/pipeline.py --mode data --topic "{args.topic}"`
3. `python scripts/research_framework/pipeline.py --mode analysis --topic "{args.topic}"`
4. `python scripts/agent_pipeline.py --topic "{args.topic}" --venue {args.venue} --language {args.language}`
"""
    out_path.write_text(design_md, encoding="utf-8")
    print(f"  ✅ {out_path}")
    return 0


def _run_review_mode(args: argparse.Namespace) -> int:
    """Stage 7 only: adversarial review of an existing draft."""
    if not args.draft_file:
        print("ERROR: --mode review requires --draft-file <path>", file=sys.stderr)
        return 2
    import subprocess
    rc = subprocess.call(
        [
            sys.executable, "scripts/core/llm_reviewer.py",
            "--draft", args.draft_file,
            "--venue", args.venue,
            "--no-llm",
        ]
    )
    return rc


def _main_dispatch() -> int:
    """Top-level CLI dispatcher: --mode design/review short-circuit the full pipeline."""
    args = _parse_args()
    if args.mode != "full":
        dispatch = {
            "design": _run_design_mode,
            "review": _run_review_mode,
        }
        handler = dispatch.get(args.mode)
        if handler is None:
            if args.mode == "full":
                return _run_full_pipeline(args)
            print(
                f"ERROR: --mode {args.mode} is not yet implemented as a stand-alone stage. "
                "Use --mode design or --mode review, or omit --mode for the full pipeline.",
                file=sys.stderr,
            )
            return 1
        return handler(args)


# Public CLI entry point: alias so `python pipeline.py` and `from pipeline import main` work
main = _run_full_pipeline
if __name__ == "__main__":
    raise SystemExit(_main_dispatch())
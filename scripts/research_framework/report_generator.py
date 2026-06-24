"""
research_framework/report_generator.py
Universal academic paper generator — outputs LaTeX and Word (.docx) with embedded tables.

Key features:
- LaTeX (.tex) with booktabs/threeparttable tables
- Word (.docx) with real python-docx Table objects (NOT images or text)
- Full Chinese and English language support
- Auto-generates all standard tables: Descriptive, DID, Heterogeneity, Mechanisms
- Includes provenance summary in appendix
- Reproduction manifest (data hash, timestamp, sources)

Usage:
    gen = ReportGenerator(output_dir="output/")
    gen.set_language("zh")       # or "en"
    gen.set_title("ESG表现与融资约束")
    gen.add_section("Introduction", text)
    gen.add_table("tab:did", table_data, caption="Table 3: DID Results")
    gen.add_figure("fig1.png", caption="Figure 1")
    gen.generate_tex("paper.tex")
    gen.generate_docx("paper.docx")
"""

from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd

_log = logging.getLogger("report_generator")
_log.setLevel(logging.INFO)

# Shared data types — single source of truth (also used by pipeline.py and data_fetcher.py)
from scripts.research_framework.base import DataSource

# ─────────────────────────────────────────
# LaTeX escaping helper (prevents injection in \includegraphics paths)
# ─────────────────────────────────────────

_LATEX_ESCAPE_TABLE = str.maketrans({
    "\\": r"\textbackslash{}",
    "{": r"\{",
    "}": r"\}",
    "$": r"\$",
    "#": r"\#",
    "%": r"\%",
    "&": r"\&",
    "_": r"\_",
    "^": r"\textasciicircum{}",
    "~": r"\textasciitilde{}",
})


def _latex_escape(s: str) -> str:
    """Escape LaTeX special characters to prevent injection."""
    return s.translate(_LATEX_ESCAPE_TABLE)
# TRANSLATION DICTIONARIES
# ─────────────────────────────────────────
ZH_EN = {
    # Paper metadata
    "title":          "标题",
    "author":         "作者",
    "date":           "日期",
    "keywords":       "关键词",
    "abstract":       "摘要",
    "see_also": "另见",
    "section":       "节",
    "table":         "表",
    "figure":        "图",
    "references":     "参考文献",
    "appendix":      "附录",
    "notes":         "注释",
    "source":        "数据来源",
    # Table elements
    "variable":       "变量",
    "mean":          "均值",
    "std":          "标准差",
    "min":           "最小值",
    "max":           "最大值",
    "n":             "观测数",
    "obs":           "观测",
    "coefficient":   "系数",
    "std_error":     "标准误",
    "p_value":       "p值",
    "t_stat":        "t统计量",
    "r_squared":     "R²",
    "firm_fe":       "公司固定效应",
    "year_fe":       "年度固定效应",
    "notes":         "注释",
    # Model terms
    "did":           "双重差分项",
    "treatment":     "处理变量",
    "post":          "政策后",
    "constant":      "常数项",
    # Diagnostics
    "data_source":   "数据来源",
    "simulated":     "模拟数据",
    "fallback":      "回退数据",
    "provenance":    "数据溯源",
    "reproducibility": "可复现性",
}

EN_TEXT = {
    "firm_fe":       "Firm FE",
    "year_fe":       "Year FE",
    "obs":           "Observations",
    "r_squared":     "R²",
    "constant":      "Constant",
    "did":           "ESG\\_high × Post",
}

# Provenance LaTeX macros to prepend to document when tracker is present
PROVENANCE_LATEX_MACROS = r"""
% Provenance tracking macros
\usepackage{xcolor}
\usepackage{ifthen}
\usepackage{forloop}
\usepackage{xparse}

% \provenance{type}{id}{description} - document a provenance node
\NewDocumentCommand{\provenance}{m m m}{%
  \textcolor{blue}{[#1: #2]} #3%
}

% \sourcedfrom{source_id}{target_id} - document data flow
\NewDocumentCommand{\sourcedfrom}{m m}{%
  \xrightarrow{\text{from #1}}#2%
}

% Provenance footnote for simulated data
\NewDocumentCommand{\simulatedfootnote}{}{%
  \footnote{\textcolor{red}{[SIMULATED]} 本数据为模拟数据，仅用于演示目的。}%
}
"""

# ─────────────────────────────────────────
# TABLE FORMATTERS
# ─────────────────────────────────────────
class TableFormatter:
    """Formats regression results for LaTeX and Word output."""

    @staticmethod
    def did_to_latex(
        results_list: list[dict],
        y_labels: list[str],
        x_vars: list[str],
        title: str = "",
        label: str = "",
        notes: str = "",
        sig_markers: str = "***,**,*,†",
        add_fallback_warning: bool = False,
        simulated_vars: set[str] | None = None,
    ) -> str:
        """Generate a publication-quality LaTeX table from DID results.

        Parameters
        ----------
        simulated_vars : set[str]
            Variable names that should be rendered in red (\\textcolor{red}{...})
            to indicate simulated/demonstration data.
        """
        sig_map = {"***": "p<0.01", "**": "p<0.05", "*": "p<0.10", "$\\dagger$": "p<0.15", "": ""}
        sim = simulated_vars or set()

        def _red(s: str) -> str:
            """Wrap simulated values in red."""
            return f"\\textcolor{{red}}{{{s}}}"

        col_count = 1 + len(y_labels)
        col_spec = "l" + "c" * len(y_labels)

        lines = [
            "\\begin{table}[htbp]",
            "  \\centering",
            f"  \\caption{{{title}}}",
            f"  \\label{{{label}}}",
            "  \\begin{threeparttable}",
            f"  \\begin{{tabular}}{{{col_spec}}}",
            "    \\toprule",
        ]

        # Header row
        header = "    \\textbf{Variable} & " + " & ".join(
            f"\\textbf{{{y}}}" for y in y_labels
        ) + " \\\\"
        lines.append(header)
        lines.append("    \\midrule")

        for var in x_vars:
            cells = [f"\\textit{{{var}}}"]
            for res in results_list:
                coefs = res.get("all_coefs", {})
                if var in coefs:
                    v = coefs[var]
                    c = v["coef"]; s = v["se"]; sig = v.get("sig", "")
                    val_str = f"${c:.4f}{sig}$ \\quad (${s:.4f}$)"
                    # Apply red color to simulated variable values
                    if var in sim:
                        val_str = _red(val_str)
                    cells.append(val_str)
                else:
                    cells.append("—")
            lines.append("    " + " & ".join(cells) + " \\\\")

        # Notes
        note_text = f"\\item Standard errors in parentheses. {sig_markers} indicate significance."
        if add_fallback_warning:
            note_text += " ⚠ WARNING: Firm FE dropped due to insufficient degrees of freedom."
        lines.extend([
            "    \\midrule",
            "    \\textbf{N} & " + " & ".join(str(r.get("n_obs", "N/A")) for r in results_list) + " \\\\",
            "    \\textbf{R$^2$} & " + " & ".join("{:.3f}".format(r.get("r_squared", 0)) for r in results_list) + " \\\\",
            "    \\bottomrule",
            "  \\end{tabular}",
            "  \\begin{tablenotes}",
            "    \\small",
            f"    \\item {note_text}",
            "  \\end{tablenotes}",
            "  \\end{threeparttable}",
            "\\end{table}",
        ])
        return "\n".join(lines)

    @staticmethod
    def descriptive_to_latex(
        df: pd.DataFrame,
        title: str = "Descriptive Statistics",
        label: str = "tab:descriptive",
        n_col: str = "N",
        stats: list[str] = ["mean", "std", "min", "p50", "max"],
    ) -> str:
        """Generate LaTeX descriptive stats table.

        Parameters
        ----------
        df : pd.DataFrame
            Expected format from pandas describe():
            - columns = variable names
            - index = stat names (count, mean, std, min, 50%, max, ...)
            If df.index contains column names instead (columns = stats), the df
            will be auto-transposed for convenience.
        """
        # Auto-detect format: if stat names are in columns, transpose
        stat_names = {"count", "mean", "std", "min", "50%", "max",
                      "25%", "75%", "N", "n"}
        if set(df.columns).issubset(stat_names) and set(df.index).issubset(stat_names):
            df = df.T
        elif not set(df.index).issubset(stat_names):
            df = df.T

        stat_map = {"mean": "均值", "std": "标准差", "min": "最小值",
                     "median": "中位数", "max": "最大值", "count": "N",
                     "50%": "中位数", "25%": "25%", "75%": "75%"}
        # Normalize stat name: "p50" → "50%" (pandas convention)
        def _norm(s: str) -> str:
            if s == "p50":
                return "50%"
            return s
        col_spec = "l" + "c" * (len(stats) + 1)

        lines = [
            "\\begin{table}[htbp]",
            "  \\centering",
            f"  \\caption{{{title}}}",
            f"  \\label{{{label}}}",
            "  \\begin{threeparttable}",
            f"  \\begin{{tabular}}{{{col_spec}}}",
            "    \\toprule",
            "    \\textbf{变量} & " + " & ".join(f"\\textbf{{{stat_map.get(s,s)}}}" for s in stats) + " \\\\",
            "    \\midrule",
        ]
        for var in df.columns:
            row = [f"\\textit{{{var}}}"]
            for s in stats:
                norm_s = _norm(s)
                val = None
                # 支持多种索引格式
                if norm_s in df.index:
                    val = df.loc[norm_s, var]
                elif s in df.index:
                    val = df.loc[s, var]
                # Guard: val must be a number, not NaN
                if val is None or (isinstance(val, float) and pd.isna(val)):
                    row.append("—")
                else:
                    try:
                        row.append(f"{float(val):.4f}")
                    except (ValueError, TypeError):
                        row.append("—")
            lines.append("    " + " & ".join(row) + " \\\\")
        lines.extend([
            "    \\bottomrule",
            "  \\end{tabular}",
            "  \\begin{tablenotes}",
            "    \\small",
            "    \\item \\textit{数据来源: Dynamic — see provenance appendix}",
            "  \\end{tablenotes}",
            "  \\end{threeparttable}",
            "\\end{table}",
        ])
        return "\n".join(lines)


# ─────────────────────────────────────────
# MAIN REPORT GENERATOR
# ─────────────────────────────────────────
class ReportGenerator:
    """
    Generates publication-ready academic papers in LaTeX and Word formats.
    
    Supports Chinese (zh) and English (en) output with proper UTF-8 encoding.
    Word output uses python-docx Table objects for real embedded tables.
    
    Args:
        output_dir: Directory to save generated files
        language: "zh" for Chinese, "en" for English
        provenance_tracker: Optional ProvenanceTracker for data provenance appendix
    """

    def __init__(
        self,
        output_dir: str | Path = "output/",
        language: str = "en",
        provenance_tracker=None,
    ):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.language = language
        self.tracker = provenance_tracker
        self._sections: list[dict] = []
        self._tables: list[dict] = []
        self._figures: list[dict] = []
        self._metadata: dict = {
            "title_en": "", "title_zh": "",
            "author": "", "date": datetime.now().strftime("%Y-%m-%d"),
            "abstract_en": "", "abstract_zh": "",
            "keywords_en": [], "keywords_zh": [],
        }

    def set_title(self, title_zh: str, title_en: str = ""):
        self._metadata["title_zh"] = title_zh
        self._metadata["title_en"] = title_en or title_zh

    def set_abstract(self, abstract_zh: str, abstract_en: str = ""):
        self._metadata["abstract_zh"] = abstract_zh
        self._metadata["abstract_en"] = abstract_en or abstract_zh

    def add_section(self, title: str, content: str, level: int = 1):
        """Add a section with title and Markdown content."""
        self._sections.append({"title": title, "content": content, "level": level})

    def add_table(
        self,
        label: str,
        data: pd.DataFrame | dict | str,
        caption_zh: str = "",
        caption_en: str = "",
        table_format: str = "did",
        notes: str = "",
        provenance: dict | None = None,
    ):
        """Add a table to the report. data can be DataFrame, dict of coefs, or LaTeX string."""
        self._tables.append({
            "label": label,
            "data": data,
            "caption_zh": caption_zh,
            "caption_en": caption_en,
            "format": table_format,
            "notes": notes,
            "provenance": provenance or {},
        })

    def add_figure(
        self,
        path: str | Path,
        caption_zh: str = "",
        caption_en: str = "",
        width: float = 0.9,
    ):
        """Add a figure to the report."""
        self._figures.append({
            "path": Path(path),
            "caption_zh": caption_zh,
            "caption_en": caption_en,
            "width": width,
        })

    # ─────────────────────────────────────
    # LATEX GENERATION
    # ─────────────────────────────────────
    def generate_tex(self, filename: str = "paper.tex") -> Path:
        """Generate LaTeX document."""
        path = self.output_dir / filename
        lines = self._build_tex_content()
        path.write_text("\n".join(lines), encoding="utf-8")
        _log.info(f"LaTeX saved: {path}")
        return path

    def _build_tex_content(self) -> list[str]:
        title = (self._metadata["title_zh"] if self.language == "zh"
                 else self._metadata["title_en"])
        abstract = (self._metadata["abstract_zh"] if self.language == "zh"
                   else self._metadata["abstract_en"])
        keywords = (self._metadata.get("keywords_zh", [])
                    if self.language == "zh"
                    else self._metadata.get("keywords_en", []))

        # Journal template system: use article as the base class for standalone generation.
        # When submitting to a specific journal, replace the .cls file:
        #   JF/JFE → jf.cls / jfe.cls,  RFS → rfs.cls,
        #   管理世界/经济研究/金融研究 → ctexart.cls
        journal = self._metadata.get("journal", "") or ""
        doc_class = "article"   # standalone-safe default

        lines = [
            f"\\documentclass[12pt,a4paper]{{{doc_class}}}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{geometry}",
            "\\geometry{margin=1in}",
            "\\usepackage{booktabs}",
            "\\usepackage{threeparttable}",
            "\\usepackage{amsmath,amssymb}",
            "\\usepackage{graphicx}",
            "\\usepackage{natbib}",
            "\\usepackage{setspace}",
            "\\doublespacing",
            "\\usepackage[colorlinks=true,linkcolor=blue,citecolor=blue,urlcolor=blue]{hyperref}",
            "\\usepackage{color}",
        ]

        # Provenance tracking macros (if tracker is available)
        if self.tracker is not None:
            lines.append(PROVENANCE_LATEX_MACROS)

        lines.extend([
            "\\begin{document}",
            f"\\title{{{title}}}",
            "\\date{\\today}",
            "\\maketitle",
            "\\begin{abstract}",
            abstract,
            "\\end{abstract}",
            "\\newpage",
        ])

        for sec in self._sections:
            level = sec["level"]
            if level == 1:
                lines.append(f"\\section{{{sec['title']}}}")
            elif level == 2:
                lines.append(f"\\subsection{{{sec['title']}}}")
            lines.append(sec["content"])
            lines.append("")

        # Tables
        for tbl in self._tables:
            cap = (tbl["caption_zh"] if self.language == "zh" else tbl["caption_en"])
            if isinstance(tbl["data"], str):
                lines.append(tbl["data"])  # Raw LaTeX
            elif isinstance(tbl["data"], dict):
                # DID results dict
                prov = tbl.get("provenance", {})
                sim_vars = {
                    k for k, v in prov.items()
                    if isinstance(v, dict) and v.get("source") == DataSource.SIMULATED
                }
                latex = TableFormatter.did_to_latex(
                    [tbl["data"]], [f"({i+1})" for i in range(1)],
                    list(tbl["data"].get("all_coefs", {}).keys())[:6],
                    title=cap, label=tbl["label"], notes=tbl["notes"],
                    add_fallback_warning=bool(sim_vars),  # 有模拟数据时显示警告
                    simulated_vars=sim_vars,
                )
                lines.append(latex)
            lines.append("")

        # Figures
        for fig in self._figures:
            cap = (fig["caption_zh"] if self.language == "zh" else fig["caption_en"])
            lines.extend([
                "\\begin{figure}[htbp]",
                "  \\centering",
                f"  \\includegraphics[width={fig['width']}\\textwidth]{{{_latex_escape(fig['path'].name)}}}",
                f"  \\caption{{{_latex_escape(cap)}}}",
                "\\end{figure}",
            ])

        # References — use bib file from output_dir if it exists; graceful fallback otherwise
        bib_path = self.output_dir / "references.bib"
        if bib_path.exists():
            bib_cmd = f"\\bibliography{{{str(bib_path)}}}"
        else:
            _log.warning(
                "references.bib not found in %s. "
                "Run fin-ref-paper skill or create references.bib before compiling.",
                self.output_dir
            )
            bib_cmd = "% \\bibliography{references}  % file not found"
        lines.extend([
            "\\newpage",
            "\\section*{References}",
            "\\bibliographystyle{plainnat}",
            bib_cmd,
        ])

        # Appendix
        lines.extend(["\\newpage", "\\appendix"])
        if self.tracker:
            lines.extend(["\\section{A. Data Provenance Summary}", self._build_provenance_appendix()])

        lines.append("\\end{document}")
        return lines

    def _build_provenance_appendix(self) -> str:
        r"""Build LaTeX provenance appendix (caller provides the \section header)."""
        if not self.tracker:
            return ""
        summary = self.tracker.summary()
        sim_fields = self.tracker.simulated_fields()
        n_fields = summary.get("total_fields", 0)
        lines = [
            "\\begin{center}",
            "  \\fbox{\\parbox{0.95\\textwidth}{",
            "    \\textbf{\\textdbend\\ 数据溯源说明} \\\\",
            "    本附录记录报告中所有数据字段的来源，确保分析结果可复现。",
            "    模拟数据（\\textit{simulated}）\\textbf{仅用于演示目的}，",
            "    不得在正式发表中使用。",
            "  }}",
            "}\\end{center}",
            "",
            f"Total fields tracked: {n_fields}",
            "\\begin{itemize}",
        ]
        for src, cnt in summary.get("by_source", {}).items():
            lines.append(f"  \\item \\textit{{{src}}}: {cnt} fields")
        if sim_fields:
            lines.extend([
                r"  \item[\textdbend]",
                r"    \textcolor{red}{\textbf{WARNING: Simulated data — DEMONSTRATION ONLY}} \\",
                f"    Simulated fields: \\texttt{{{', '.join(sim_fields)}}}",
                r"    These values were synthetically generated and are \textbf{not} derived from real data.",
                r"    Do \textbf{not} use in published work.",
            ])
        lines.extend(["\\end{itemize}"])
        return "\n".join(lines)

    # ─────────────────────────────────────
    # WORD (DOCX) GENERATION
    # ─────────────────────────────────────
    def generate_docx(self, filename: str = "paper.docx") -> Path | None:
        """Generate Word document with REAL embedded tables (not images)."""
        import importlib.util

        if importlib.util.find_spec("docx") is None:
            _log.error("python-docx not installed. Run: pip install python-docx")
            return None

        from docx import Document as DocxDocument
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor

        doc = DocxDocument()
        self._apply_docx_styles(doc)

        title = (self._metadata["title_zh"] if self.language == "zh"
                 else self._metadata["title_en"])
        abstract = (self._metadata["abstract_zh"] if self.language == "zh"
                   else self._metadata["abstract_en"])

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        for k, v in [("Author", self._metadata.get("author", "")),
                      ("Date", self._metadata["date"])]:
            if v:
                p = doc.add_paragraph()
                p.add_run(f"{k}: ").bold = True
                p.add_run(v)

        # Abstract
        doc.add_heading(
            "摘要 / Abstract" if self.language == "zh" else "Abstract",
            level=1
        )
        doc.add_paragraph(abstract)

        # Sections
        for sec in self._sections:
            doc.add_heading(sec["title"], level=sec.get("level", 1) + 1)
            doc.add_paragraph(sec["content"])

        # ── REAL EMBEDDED TABLES ──
        for tbl in self._tables:
            self._add_docx_table(doc, tbl)

        # Figures
        for fig in self._figures:
            cap = (fig["caption_zh"] if self.language == "zh" else fig["caption_en"])
            self._add_docx_figure(doc, fig["path"], cap)

        # Provenance appendix
        if self.tracker:
            doc.add_page_break()
            self._add_docx_provenance(doc)

        path = self.output_dir / filename
        doc.save(str(path))
        _log.info(f"Word document saved: {path}")
        return path

    def _add_docx_table(self, doc, tbl: dict):
        """Add a REAL python-docx Table (not image) to the document."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        cap = (tbl["caption_zh"] if self.language == "zh" else tbl["caption_en"])

        # Table caption
        p = doc.add_paragraph()
        run = p.add_run(f"{cap}")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        data = tbl["data"]

        if isinstance(data, pd.DataFrame):
            df = data
        elif isinstance(data, dict):
            # Convert results dict to DataFrame
            coefs = data.get("all_coefs", {})
            if not coefs:
                doc.add_paragraph("[Table data unavailable]")
                return
            rows_data = []
            for var, v in coefs.items():
                sig = v.get("sig", "")
                rows_data.append([
                    var,
                    f"{v['coef']:.4f}{sig}",
                    f"({v['se']:.4f})",
                    f"{v['pval']:.4f}",
                ])
            df = pd.DataFrame(rows_data, columns=["Variable", "Coef", "SE", "p-value"])
        elif isinstance(data, str):
            # LaTeX fallback — add as text
            doc.add_paragraph(f"[LaTeX table: {tbl['label']}]")
            return
        else:
            doc.add_paragraph("[Unsupported table format]")
            return

        if df.empty:
            doc.add_paragraph("[Empty table]")
            return

        # ── CREATE REAL DOCX TABLE ──
        n_cols = len(df.columns)
        n_rows = len(df) + 1  # +1 for header

        tbl_obj = doc.add_table(rows=n_rows, cols=n_cols)
        tbl_obj.style = "Table Grid"
        tbl_obj.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = tbl_obj.rows[0]
        for j, col_name in enumerate(df.columns):
            cell = hdr.cells[j]
            cell.text = str(col_name)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        #         # Data rows — red for simulated fields (provenance-driven).
        # Provenance keys use formats like "{field}" or "{ticker}:{year}:{field}".
        # Extract field name (last component) for matching with regression variable names.
        provenance = tbl.get("provenance", {})
        simulated_vars = set()
        for k, v in provenance.items():
            if isinstance(v, dict) and v.get("source") == DataSource.SIMULATED:
                field_part = k.rsplit(":", 1)[-1]
                simulated_vars.add(field_part)
        RED = RGBColor(0xC0, 0x00, 0x00)  # dark red for simulated data

        for i, (_, row) in enumerate(df.iterrows()):
            cells = tbl_obj.rows[i + 1].cells
            var_name = str(row.get(df.columns[0], ""))
            is_simulated = var_name in simulated_vars
            for j, val in enumerate(row):
                cells[j].text = str(val)
                for para in cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
                        # Apply red color for simulated variable rows
                        if is_simulated:
                            run.font.color.rgb = RED
                            run.bold = True
                        # Bold the DID coefficient row
                        if var_name in ["did", "psm_did", "ESG_high × Post"]:
                            run.bold = True

        # Notes
        if tbl.get("notes"):
            p = doc.add_paragraph()
            p.add_run(f"注: {tbl['notes']}").italic = True
            for run in p.runs:
                run.font.size = Pt(8)

        doc.add_paragraph()  # Spacing

    def _add_docx_figure(self, doc, path: Path, caption: str):
        """Add a figure image to the Word document."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        if not path.exists():
            doc.add_paragraph(f"[Figure not found: {path}]")
            return
        try:
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(str(path), width=Inches(5.5))
            p = doc.add_paragraph()
            p.add_run(caption).italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            _log.warning(f"Could not embed figure {path}: {e}")
            doc.add_paragraph(f"[Figure: {caption}]")

    def _add_docx_provenance(self, doc):
        """Add data provenance appendix to Word document."""
        from docx.shared import RGBColor
        doc.add_heading("附录A: 数据溯源 / Appendix A: Data Provenance", level=1)

        # ── WARNING BANNER (red) ──
        p = doc.add_paragraph()
        run = p.add_run(
            "\u26a0 数据溯源说明 / Data Provenance Notice\n"
            "本附录记录所有数据来源，确保可复现性。\n"
            "标注为 SIMULATED 的字段仅用于演示，不得在正式发表中使用。"
        )
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.alignment = None  # left-aligned

        summary = self.tracker.summary()
        doc.add_paragraph(f"总字段数 / Total fields tracked: {summary.get('total_fields', 0)}")
        for src, cnt in summary.get("by_source", {}).items():
            doc.add_paragraph(f"\u2022 {src}: {cnt} 字段/fields")
        sim = self.tracker.simulated_fields()
        if sim:
            p = doc.add_paragraph()
            run = p.add_run(
                f"\u26a0 WARNING: Simulated data detected — FOR DEMONSTRATION ONLY\n"
                f"Simulated fields: {', '.join(sim)}\n"
                f"这些值为合成生成，非真实数据。"
            )
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    def _apply_docx_styles(self, doc):
        """Set document-wide styles."""
        try:
            from docx.shared import Pt
            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(11)
        except Exception:  # noqa: S110
            pass

    # ─────────────────────────────────────
    # HELPERS
    # ─────────────────────────────────────
    def set_language(self, lang: str):
        self.language = lang

    def save_manifest(self, extra: dict | None = None):
        manifest = {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "language": self.language,
            "title_zh": self._metadata["title_zh"],
            "title_en": self._metadata["title_en"],
            "n_sections": len(self._sections),
            "n_tables": len(self._tables),
            "n_figures": len(self._figures),
            "provenance_summary": self.tracker.summary() if self.tracker else {},
        }
        if extra:
            manifest.update(extra)
        path = self.output_dir / "manifest.json"
        path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False, default=str))
        _log.info(f"Manifest saved: {path}")

    # ─────────────────────────────────────
    # PAPER GENERATION (P0-1: end-to-end PDF)
    # ─────────────────────────────────────
    def generate_paper(
        self,
        topic: str,
        outline: dict,
        data: pd.DataFrame | None = None,
        regressions: dict | None = None,
        references: list | None = None,
        journal: str = "经济研究",
        output_dir: str | Path | None = None,
    ) -> Path:
        """Generate a complete paper as .tex and .pdf.

        This is the P0-1 entry point that wires the entire research pipeline
        output into a submission-ready PDF.

        Args:
            topic: Research topic (used as paper title if no title in outline)
            outline: Paper outline dict with keys:
                abstract, intro/introduction, lit_review, method/methodology,
                results, robustness/robustness_checks, conclusion
            data: Optional DataFrame for reproducibility section
            regressions: Dict of regression results keyed by table name
            references: List of BibTeX entries (as str or dict)
            journal: Target journal name (e.g. "经济研究", "JFE")
            output_dir: Output directory (default: self.output_dir)

        Returns:
            Path to the generated .tex file
        """
        import logging as _rg_log
        _rg_log = _rg_log.getLogger("report_generator.generate_paper")

        # ── Output directory ───────────────────────────────────────────────────────
        out_dir = Path(output_dir) if output_dir else self.output_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        # Determine language from journal
        is_chinese = any(j in journal for j in [
            "经济研究", "金融研究", "管理世界", "中国", "会计研究",
            "财政研究", "世界经济", "统计研究", "数量经济",
        ])
        lang = "zh" if is_chinese else self.language
        self.set_language(lang)

        # ── Title ────────────────────────────────────────────────────────────────
        title_zh = outline.get("title_zh") or outline.get("title") or topic
        title_en = outline.get("title_en", title_zh)
        self.set_title(title_zh, title_en)

        # ── Abstract ────────────────────────────────────────────────────────────
        abstract_zh = outline.get("abstract_zh") or outline.get("abstract", "")
        abstract_en = outline.get("abstract_en", abstract_zh)
        self.set_abstract(abstract_zh, abstract_en)

        # ── Keywords ─────────────────────────────────────────────────────────────
        if "keywords_zh" in outline:
            self._metadata["keywords_zh"] = outline["keywords_zh"]
        if "keywords_en" in outline:
            self._metadata["keywords_en"] = outline["keywords_en"]

        # ── Sections ─────────────────────────────────────────────────────────────
        section_order = [
            ("introduction", "Introduction"),
            ("intro", "Introduction"),
            ("lit_review", "Literature Review"),
            ("literature_review", "Literature Review"),
            ("method", "Methodology"),
            ("methodology", "Methodology"),
            ("data", "Data"),
            ("results", "Results"),
            ("main_results", "Main Results"),
            ("heterogeneity", "Heterogeneity Analysis"),
            ("mechanism", "Mechanism Analysis"),
            ("robustness", "Robustness Checks"),
            ("robustness_checks", "Robustness Checks"),
            ("conclusion", "Conclusion"),
            ("conclusions", "Conclusion"),
        ]
        processed_keys = set()
        for key, label in section_order:
            if key in outline and key not in processed_keys:
                content = outline[key]
                if isinstance(content, str):
                    self.add_section(label, content, level=1)
                elif isinstance(content, dict):
                    title = content.get("title", label)
                    body = content.get("content", content.get("text", ""))
                    self.add_section(title, body, level=1)
                processed_keys.add(key)

        # Any remaining string keys in outline not yet processed → add as-is
        for key, val in outline.items():
            if key in processed_keys or not isinstance(val, str):
                continue
            self.add_section(key.replace("_", " ").title(), val, level=1)

        # ── Regression tables ───────────────────────────────────────────────────
        if regressions:
            for tbl_name, tbl_data in regressions.items():
                if isinstance(tbl_data, str):
                    self.add_table(f"tab:{tbl_name}", tbl_data,
                                   caption_zh=tbl_name, table_format="raw")
                elif isinstance(tbl_data, dict):
                    self.add_table(f"tab:{tbl_name}", tbl_data,
                                   caption_zh=tbl_name)
                elif isinstance(tbl_data, pd.DataFrame):
                    self.add_table(f"tab:{tbl_name}", tbl_data,
                                   caption_zh=tbl_name, table_format="descriptive")

        # ── References ──────────────────────────────────────────────────────────
        if references:
            bib_path = out_dir / "references.bib"
            bib_entries = "\n\n".join(
                r if isinstance(r, str) else r.get("bibtex", str(r))
                for r in references
            )
            bib_path.write_text(bib_entries, encoding="utf-8")
            _rg_log.info("References written: %s", bib_path)

        # ── Generate LaTeX ─────────────────────────────────────────────────────
        tex_filename = self._sanitize_filename(topic) + ".tex"
        tex_path = out_dir / tex_filename
        gen_path = self.generate_tex(tex_filename)
        _rg_log.info("LaTeX generated: %s", gen_path)

        # ── Compile to PDF ─────────────────────────────────────────────────────
        # Import journal_template lazily (avoids circular dep at module level)
        try:
            from scripts.journal_template import get_template
        except ImportError:
            _rg_log.warning("scripts.journal_template not available, skipping PDF compile")
            return tex_path

        template = get_template(journal)
        if template is None:
            _rg_log.warning("Unknown journal %r, using xelatex directly", journal)
            template = get_template("JFE")  # fallback to JFE

        engine = "xelatex" if is_chinese else "pdflatex"
        _rg_log.info("Compiling %s with engine=%s", tex_path, engine)

        try:
            success = template.compile(str(tex_path), engine=engine, passes=2)
            if success:
                _rg_log.info("PDF compiled successfully: %s", tex_path.with_suffix(".pdf"))
            else:
                _rg_log.warning("PDF compilation returned False (check TeX installation)")
        except Exception as e:
            _rg_log.warning("PDF compilation failed: %s", e)

        # ── Save manifest ──────────────────────────────────────────────────────
        self.save_manifest({
            "topic": topic,
            "journal": journal,
            "tex_path": str(tex_path),
            "output_dir": str(out_dir),
        })

        return tex_path

    @staticmethod
    def _sanitize_filename(topic: str) -> str:
        """Convert a topic string into a safe filename."""
        import re
        safe = re.sub(r"[^\w\u4e00-\u9fff\- ]", "_", topic)
        safe = re.sub(r"[\s]+", "_", safe.strip())
        return safe[:80] or "paper"


__all__ = ["ReportGenerator", "TableFormatter"]


# ─────────────────────────────────────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(
        prog="report_generator.py",
        description="经济金融论文生成器 — 输出 LaTeX 和 Word 格式（含数据溯源）。",
    )
    parser.add_argument("--topic", "-t", required=True, help="论文主题")
    parser.add_argument(
        "--outline", "-o",
        help="大纲文件路径（JSON 或 YAML），或 PAPER_OUTLINE.md",
    )
    parser.add_argument(
        "--journal", "-j",
        default="JFE",
        help="期刊模板 (JFE/JF/RFS/经济研究/金融研究, default: JFE)",
    )
    parser.add_argument(
        "--language", "-l",
        default="zh",
        choices=["zh", "en"],
        help="论文语言 (zh/en, default: zh)",
    )
    parser.add_argument(
        "--output-dir", "-d",
        default="output/fin-manuscript",
        help="输出目录 (default: output/fin-manuscript)",
    )
    parser.add_argument(
        "--no-compile",
        action="store_true",
        help="仅生成 LaTeX，不编译 PDF",
    )

    args = parser.parse_args()

    from pathlib import Path

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # 读取大纲
    outline = {}
    if args.outline:
        p = Path(args.outline)
        if p.exists():
            raw = p.read_text(encoding="utf-8")
            if p.suffix in (".json",):
                import json
                outline = json.loads(raw)
            else:
                import yaml
                outline = yaml.safe_load(raw) or {}
            print(f"大纲已加载: {p.name} ({len(outline)} 个章节)")
        else:
            print(f"警告: 大纲文件不存在: {args.outline}，使用空大纲", file=sys.stderr)

    print(f"生成论文: {args.topic}")
    print(f"期刊: {args.journal} | 语言: {args.language} | 输出: {out_dir}")

    gen = ReportGenerator(output_dir=str(out_dir))
    gen.set_language(args.language)

    try:
        tex_path = gen.generate_paper(
            topic=args.topic,
            outline=outline,
            regressions={},
            references=[],
            journal=args.journal,
        )
        print(f"\nLaTeX 已生成: {tex_path}")
        if not args.no_compile:
            print("PDF 编译（如需）：xelatex/pdflatex {tex_path}")
    except Exception as e:
        print(f"生成失败: {e}", file=sys.stderr)
        sys.exit(1)


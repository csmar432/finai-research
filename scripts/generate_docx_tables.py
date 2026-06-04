#!/usr/bin/env python3
"""生成规范三线表Word文档"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

# ── 三线表格式工具 ──────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), v)
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_three_line(table):
    """
    将表格设为标准三线表格式（优化版）：
    - 顶线：1.5pt 黑色实线
    - 底线：1.5pt 黑色实线
    - 栏目线（表头下）：0.75pt 黑色实线
    - 中间行：仅上下细线，左右无线（学术规范）
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 表格居中对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)

    # 使用表格级边框设置（更高效）
    tblBorders = OxmlElement('w:tblBorders')
    # 顶部粗线
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)
    # 底部粗线
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)
    # 左边线
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '4')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '000000')
    tblBorders.append(left)
    # 右边线
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), '000000')
    tblBorders.append(right)
    # 内部水平线（栏目线）
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '6')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), '000000')
    tblBorders.append(insideH)
    # 内部垂直线（无线）
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'single')
    insideV.set(qn('w:sz'), '0')
    insideV.set(qn('w:space'), '0')
    insideV.set(qn('w:color'), 'FFFFFF')
    tblBorders.append(insideV)
    tblPr.append(tblBorders)

    # 仅对表头行设置加粗和居中（避免逐单元格设置边框）
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            # 表头加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            # 表头居中
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_text(cell, text, bold=False, size=10, center=False):
    """设置单元格文本格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_header_style(cell, text, size=10):
    """表头单元格样式（加粗居中）"""
    set_cell_text(cell, text, bold=True, size=size, center=True)


# ── 公式解析 ───────────────────────────────────────────────────

def parse_formula_block(lines: list[str], start_idx: int) -> tuple[str, int]:
    """
    解析一个完整的公式块（支持单行/多行 $$...$$）。

    支持格式：
      $$formula$$
      $$
      multi-line formula
      $$
      $$ inline $$ in paragraph

    Returns:
        (formula_text, next_idx): 公式文本和下一个待处理行索引
    """
    if start_idx >= len(lines):
        return "", start_idx

    line = lines[start_idx]
    next_idx = start_idx + 1

    # 如果本行没有 $$，返回空
    if '$$' not in line:
        return "", start_idx

    count = line.count('$$')

    if count >= 2:
        # 单行公式：$$...$$
        if not line.strip().startswith('$$'):
            # 行内混合：text $$formula$$ text
            # 找出第一个 $$
            idx1 = line.index('$$')
            idx2 = line.index('$$', idx1 + 2)
            content = line[idx1 + 2:idx2]
            return content.strip(), start_idx

        # 整行公式：$$...$$
        idx1 = line.index('$$')
        idx2 = line.index('$$', idx1 + 2)
        content = line[idx1 + 2:idx2]
        return content.strip(), start_idx

    # count == 1: 可能是开始或结束标记
    stripped = line.strip()

    if stripped.startswith('$$'):
        # 开始标记 $$
        idx = line.index('$$')
        before = line[:idx]
        collected = []
        if before.strip():
            collected.append(before)

        # 收集后续行直到遇到下一个 $$
        for li in range(start_idx + 1, len(lines)):
            l = lines[li]
            if '$$' in l:
                idx = l.index('$$')
                after = l[:idx]
                if after.strip():
                    collected.append(after)
                next_idx = li + 1
                break
            collected.append(l)
        else:
            next_idx = len(lines)

        return '\n'.join(collected).strip(), next_idx

    if stripped.endswith('$$'):
        # 结束标记 $$
        idx = line.index('$$')
        before = line[:idx]
        return before.strip(), start_idx + 1

    return "", start_idx


def parse_latex_env(lines: list[str], start_idx: int) -> tuple[str, int]:
    r"""解析 LaTeX 环境（\begin...\end）。"""
    if start_idx >= len(lines) or '\\begin{' not in lines[start_idx]:
        return "", start_idx

    line = lines[start_idx]
    collected = [line]
    brace_depth = line.count('{') - line.count('}')
    next_idx = start_idx + 1

    while brace_depth > 0 and next_idx < len(lines):
        l = lines[next_idx]
        collected.append(l)
        brace_depth += l.count('{') - l.count('}')
        next_idx += 1

    return '\n'.join(collected).strip(), next_idx


# ── 主转换函数 ─────────────────────────────────────────────────

def parse_markdown_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """解析 Markdown 表格，返回 (headers, rows_data)"""
    data_lines = [l for l in lines if l.strip().startswith('|') and '---' not in l]
    if not data_lines:
        return [], []

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows[0], rows[1:]


def md_to_docx(paper_text: str, output_path: str):
    """将 Markdown 论文转换为 Word 文档（保留三线表和公式）。"""
    doc = Document()

    # ── 页面设置 ──────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 默认样式 ──────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 辅助函数 ─────────────────────────────────────────
    def add_title(text: str, level: str = 'main'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        if level == 'main':
            run.font.size = Pt(18)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 'sub':
            run.font.size = Pt(14)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 'section':
            run.font.size = Pt(14)
            run.bold = True
        elif level == 'subsection':
            run.font.size = Pt(12)
            run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_para(text: str, indent: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(24)
        return p

    def add_formula(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_table_note(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(3)
        return p

    def insert_table(headers: list[str], rows_data: list[list[str]]):
        """插入三线表"""
        if not rows_data:
            return
        n_cols = len(headers)
        table = doc.add_table(rows=len(rows_data) + 1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            set_header_style(table.cell(0, j), h)

        for i, row_data in enumerate(rows_data):
            for j, val in enumerate(row_data):
                if j < n_cols:
                    center = (j > 0)
                    set_cell_text(table.cell(i + 1, j), val, center=center)

        set_table_three_line(table)
        return table

    def insert_note(note_text: str):
        """插入表格注释"""
        p = doc.add_paragraph()
        run = p.add_run(note_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(12)

    def process_inline_formula(text: str) -> "Paragraph":
        """处理行内公式 $...$"""
        p = doc.add_paragraph()
        # 用正则分割：普通文本 和 $...$ 公式
        parts = re.split(r'(\$[^$\n]+\$)', text)
        for part in parts:
            if re.match(r'\$[^$\n]+\$', part):
                r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.italic = True
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(24)
        return p

    def process_mixed_block(text: str) -> "Paragraph":
        """处理包含 $$...$$ 的段落"""
        p = doc.add_paragraph()
        # 分割普通文本和 $$公式$$
        parts = re.split(r'(\$\$[^$]+\$\$)', text)
        for part in parts:
            if re.match(r'\$\$[^$]+\$\$', part):
                formula = part[2:-2].strip()
                r = p.add_run(formula)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.italic = True
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(24)
        return p

    # ── 逐行解析 ─────────────────────────────────────────
    lines = paper_text.split('\n')
    i = 0
    pending_table: list[str] = []

    while i < len(lines):
        line = lines[i].strip()

        # ── 检测 Markdown 表格 ──
        if line.startswith('|'):
            pending_table.append(lines[i])
            i += 1
            continue

        # ── 处理累积的表格 ──
        if pending_table:
            headers, rows_data = parse_markdown_table(pending_table)
            if headers and rows_data:
                insert_table(headers, rows_data)
                # 表格注释
                if i < len(lines) and lines[i].strip().startswith('*注：'):
                    note_text = lines[i].strip().lstrip('*')
                    insert_note(note_text)
                    i += 1
                else:
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                    if i < len(lines) and lines[i].strip().startswith('*注：'):
                        note_text = lines[i].strip().lstrip('*')
                        insert_note(note_text)
                        i += 1
                    else:
                        doc.add_paragraph()
            pending_table = []
            continue

        # ── 表格注释 ──
        if line.startswith('*注：'):
            insert_note(line.lstrip('*'))
            i += 1
            continue

        # ── 分隔线 ──
        if line == '---':
            i += 1
            continue

        # ── 公式块（$$ 开头） ──
        if line.startswith('$$'):
            formula_text, next_i = parse_formula_block(lines, i)
            if formula_text:
                add_formula(formula_text)
            i = next_i
            continue

        # ── LaTeX 环境 ──
        if '\\begin{' in line:
            env_text, next_i = parse_latex_env(lines, i)
            if env_text:
                add_formula(env_text)
            i = next_i
            continue

        # ── 标题 ──
        if line.startswith('# '):
            text = line[2:].strip()
            if text:
                add_title(text, 'main')
        elif line.startswith('## ') and not line[3:].strip().startswith('**'):
            text = line[3:].strip()
            if text and not text.startswith('表'):
                add_title(text, 'sub')
        elif line.startswith('### '):
            text = line[4:].strip()
            if text:
                if text.startswith('表'):
                    i += 1
                    continue
                add_title(text, 'section')
        elif line.startswith('#### '):
            text = line[5:].strip()
            if text:
                add_title(text, 'subsection')
        elif line.startswith('**') and line.endswith('**') and '表' in line:
            pass  # 表格标题已在上方处理

        # ── 列表 ──
        elif line.startswith('- '):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
            add_bullet(text)
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)

        # ── 参考文献标题 ──
        elif line.startswith('### 参考文献') or line == '参考文献':
            p = doc.add_paragraph()
            run = p.add_run('参考文献')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # ── 参考文献条目 ──
        elif re.match(r'^\d+\.', line) and '(' in line:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            run = p.add_run(line[line.index('.') + 1:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(4)

        # ── 段落（行内公式或普通文本） ──
        elif line:
            has_inline = '$' in line and '$$' not in line
            has_block = '$$' in line and not line.startswith('$$')
            if has_inline:
                process_inline_formula(line)
            elif has_block:
                process_mixed_block(line)
            else:
                add_para(line)

        i += 1

    # 处理末尾可能残留的表格
    if pending_table:
        headers, rows_data = parse_markdown_table(pending_table)
        if headers and rows_data:
            insert_table(headers, rows_data)

    doc.save(output_path)
    print(f"✅ Word文档已保存: {output_path}")


# ── 主程序 ─────────────────────────────────────────────────
if __name__ == "__main__":
    from pathlib import Path

    matching_files = sorted(
        (Path(__file__).parent.parent / "output").glob('关税政策影响研究_去AI版_*.md')
    )
    if not matching_files:
        print("❌ 未找到匹配的论文文件: output/关税政策影响研究_去AI版_*.md")
        print("   请先生成去AI润色版论文，或修改文件名匹配模式。")
        import sys
        sys.exit(1)

    paper_file = matching_files[-1]
    paper_text = paper_file.read_text(encoding='utf-8')

    output_file = Path(__file__).parent.parent / "output/关税政策影响研究_三线表版.docx"
    md_to_docx(paper_text, str(output_file))

    print(f"\n生成完成！文件: {output_file}")
    print(f"论文字数: {len(paper_text)} 字")
